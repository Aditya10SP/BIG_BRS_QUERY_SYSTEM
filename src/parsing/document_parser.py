"""Document parser for .docx and .pdf files."""

import logging
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Dict, List

import docx
import pdfplumber
from docx.document import Document as DocxDocument
from docx.text.paragraph import Paragraph

from src.utils.errors import ParsingError, ErrorContext, log_error_with_context


logger = logging.getLogger(__name__)


@dataclass
class Section:
    """Represents a document section with hierarchical structure."""
    section_id: str
    heading: str
    level: int  # Heading level (1, 2, 3...)
    text: str
    page_numbers: List[int] = field(default_factory=list)


@dataclass
class ParsedDocument:
    """Represents a parsed document with structured content."""
    doc_id: str
    title: str
    sections: List[Section]
    metadata: Dict[str, Any] = field(default_factory=dict)


class DocumentParser:
    """Parser for extracting text and structure from .docx and .pdf files."""
    
    SUPPORTED_FORMATS = {".docx", ".pdf"}
    
    def parse(self, file_path: str, file_type: str) -> ParsedDocument:
        """
        Parse document and extract structured content.
        
        Args:
            file_path: Path to document file
            file_type: 'docx' or 'pdf'
            
        Returns:
            ParsedDocument with text, sections, metadata
            
        Raises:
            ParsingError: If document cannot be parsed
        """
        # Validate file type
        file_type = file_type.lower()
        if f".{file_type}" not in self.SUPPORTED_FORMATS:
            context = ErrorContext().add("file_type", file_type).add("file_path", file_path)
            raise ParsingError(
                f"Unsupported file format: {file_type}. "
                f"Supported formats: {', '.join(self.SUPPORTED_FORMATS)}",
                context=context.build()
            )
        
        # Validate file exists
        path = Path(file_path)
        if not path.exists():
            context = ErrorContext().add("file_path", file_path)
            raise ParsingError(f"File not found: {file_path}", context=context.build())
        
        if not path.is_file():
            context = ErrorContext().add("file_path", file_path)
            raise ParsingError(f"Path is not a file: {file_path}", context=context.build())
        
        # Parse based on file type
        try:
            if file_type == "docx":
                return self._parse_docx(file_path)
            elif file_type == "pdf":
                return self._parse_pdf(file_path)
            else:
                context = ErrorContext().add("file_type", file_type).add("file_path", file_path)
                raise ParsingError(f"Unsupported file type: {file_type}", context=context.build())
        except ParsingError:
            raise
        except Exception as e:
            log_error_with_context(
                e,
                component="DocumentParser",
                operation="parse",
                file_path=file_path,
                file_type=file_type
            )
            context = ErrorContext().add("file_path", file_path).add("file_type", file_type)
            raise ParsingError(
                f"Failed to parse {file_type} file '{path.name}': {str(e)}",
                context=context.build(),
                cause=e
            ) from e
    
    def _parse_docx(self, file_path: str) -> ParsedDocument:
        """Parse .docx file and extract structure."""
        try:
            doc = docx.Document(file_path)
        except Exception as e:
            context = ErrorContext().add("file_path", file_path).add("file_type", "docx")
            raise ParsingError(
                f"Corrupted or invalid .docx file: {str(e)}",
                context=context.build(),
                cause=e
            ) from e
        
        path = Path(file_path)
        doc_id = path.stem
        
        # Extract title (first heading or filename)
        title = self._extract_docx_title(doc) or path.stem
        
        # Extract sections with hierarchy
        sections = self._extract_docx_sections(doc)
        
        # Build metadata
        metadata = {
            "file_name": path.name,
            "file_type": "docx",
            "num_paragraphs": len(doc.paragraphs),
            "num_sections": len(sections),
        }
        
        # Add core properties if available
        if hasattr(doc, "core_properties"):
            props = doc.core_properties
            if props.author:
                metadata["author"] = props.author
            if props.created:
                metadata["created"] = props.created.isoformat()
            if props.modified:
                metadata["modified"] = props.modified.isoformat()
        
        return ParsedDocument(
            doc_id=doc_id,
            title=title,
            sections=sections,
            metadata=metadata
        )
    
    def _extract_docx_title(self, doc: DocxDocument) -> str:
        """Extract document title from first heading or paragraph."""
        for para in doc.paragraphs:
            if para.style.name.startswith("Heading 1") or para.style.name == "Title":
                text = para.text.strip()
                if text:
                    return text
        
        # Fallback to first non-empty paragraph
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                return text[:100]  # Limit title length
        
        return ""
    
    def _extract_docx_sections(self, doc: DocxDocument) -> List[Section]:
        """Extract sections with hierarchical structure from .docx."""
        sections = []
        current_section = None
        section_counter = 0
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            # Check if this is a heading
            heading_level = self._get_heading_level(para)
            
            if heading_level > 0:
                # Save previous section if exists
                if current_section is not None:
                    sections.append(current_section)
                
                # Start new section
                section_counter += 1
                current_section = Section(
                    section_id=f"s{section_counter}",
                    heading=text,
                    level=heading_level,
                    text="",
                    page_numbers=[]
                )
            else:
                # Add text to current section
                if current_section is None:
                    # Create default section for content before first heading
                    section_counter += 1
                    current_section = Section(
                        section_id=f"s{section_counter}",
                        heading="Introduction",
                        level=1,
                        text="",
                        page_numbers=[]
                    )
                
                if current_section.text:
                    current_section.text += "\n"
                current_section.text += text
        
        # Add final section
        if current_section is not None:
            sections.append(current_section)
        
        # If no sections found, create a single section with all content
        if not sections:
            all_text = "\n".join(p.text.strip() for p in doc.paragraphs if p.text.strip())
            if all_text:
                sections.append(Section(
                    section_id="s1",
                    heading="Document Content",
                    level=1,
                    text=all_text,
                    page_numbers=[]
                ))
        
        return sections
    
    def _get_heading_level(self, para: Paragraph) -> int:
        """Get heading level from paragraph style (0 if not a heading)."""
        style_name = para.style.name
        
        if style_name.startswith("Heading"):
            try:
                # Extract number from "Heading 1", "Heading 2", etc.
                level_str = style_name.replace("Heading", "").strip()
                return int(level_str)
            except (ValueError, AttributeError):
                return 0
        elif style_name == "Title":
            return 1
        
        return 0
    
    def _parse_pdf(self, file_path: str) -> ParsedDocument:
        """Parse .pdf file and extract structure."""
        try:
            pdf = pdfplumber.open(file_path)
        except Exception as e:
            context = ErrorContext().add("file_path", file_path).add("file_type", "pdf")
            raise ParsingError(
                f"Corrupted or invalid .pdf file: {str(e)}",
                context=context.build(),
                cause=e
            ) from e
        
        path = Path(file_path)
        doc_id = path.stem
        
        try:
            # Extract text from all pages
            sections = self._extract_pdf_sections(pdf)
            
            # Extract title (from metadata or first line)
            title = self._extract_pdf_title(pdf, sections) or path.stem
            
            # Build metadata
            metadata = {
                "file_name": path.name,
                "file_type": "pdf",
                "num_pages": len(pdf.pages),
                "num_sections": len(sections),
            }
            
            # Add PDF metadata if available
            if pdf.metadata:
                if pdf.metadata.get("Title"):
                    metadata["pdf_title"] = pdf.metadata["Title"]
                if pdf.metadata.get("Author"):
                    metadata["author"] = pdf.metadata["Author"]
                if pdf.metadata.get("CreationDate"):
                    metadata["created"] = pdf.metadata["CreationDate"]
            
            return ParsedDocument(
                doc_id=doc_id,
                title=title,
                sections=sections,
                metadata=metadata
            )
        finally:
            pdf.close()
    
    def _extract_pdf_title(self, pdf, sections: List[Section]) -> str:
        """Extract title from PDF metadata or first section."""
        # Try metadata first
        if pdf.metadata and pdf.metadata.get("Title"):
            title = pdf.metadata["Title"].strip()
            if title:
                return title
        
        # Try first section heading
        if sections and sections[0].heading != "Page 1":
            return sections[0].heading
        
        # Try first line of text
        if sections and sections[0].text:
            first_line = sections[0].text.split("\n")[0].strip()
            if first_line:
                return first_line[:100]
        
        return ""
    
    def _extract_pdf_sections(self, pdf) -> List[Section]:
        """Extract sections from PDF pages."""
        sections = []
        
        for page_num, page in enumerate(pdf.pages, start=1):
            try:
                text = page.extract_text()
                if not text or not text.strip():
                    continue
                
                # Create section per page (simple approach)
                # More sophisticated heading detection could be added later
                section = Section(
                    section_id=f"s{page_num}",
                    heading=f"Page {page_num}",
                    level=1,
                    text=text.strip(),
                    page_numbers=[page_num]
                )
                sections.append(section)
            except Exception as e:
                logger.warning(f"Failed to extract text from page {page_num}: {e}")
                continue
        
        if not sections:
            context = ErrorContext().add("file_path", file_path).add("num_pages", len(pdf.pages))
            raise ParsingError(
                "No text content could be extracted from PDF",
                context=context.build()
            )
        
        return sections
