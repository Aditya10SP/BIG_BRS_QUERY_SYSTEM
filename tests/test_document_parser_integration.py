"""Integration tests for DocumentParser with realistic documents."""

import pytest
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from io import BytesIO

from src.parsing import DocumentParser, ParsingError


@pytest.fixture
def parser():
    """Create DocumentParser instance."""
    return DocumentParser()


@pytest.fixture
def banking_docx(tmp_path):
    """Create a realistic banking document in .docx format."""
    file_path = tmp_path / "banking_spec.docx"
    doc = Document()
    
    # Title
    doc.add_heading("NEFT Payment System Specification", level=1)
    doc.add_paragraph("Version 2.1 - Last Updated: January 2024")
    
    # Overview section
    doc.add_heading("1. Overview", level=2)
    doc.add_paragraph(
        "The National Electronic Funds Transfer (NEFT) system enables "
        "one-to-one funds transfer between bank accounts across India."
    )
    
    # Technical specs
    doc.add_heading("2. Technical Specifications", level=2)
    doc.add_heading("2.1 Transaction Limits", level=3)
    doc.add_paragraph("Minimum: Rs. 1")
    doc.add_paragraph("Maximum: Rs. 2,00,000 per transaction")
    
    doc.add_heading("2.2 Processing Time", level=3)
    doc.add_paragraph("Batch processing in hourly settlements")
    
    # Dependencies
    doc.add_heading("3. System Dependencies", level=2)
    doc.add_paragraph("NEFT depends on:")
    doc.add_paragraph("- Core Banking System (CBS)")
    doc.add_paragraph("- Payment Gateway")
    doc.add_paragraph("- Risk Management System")
    
    doc.save(str(file_path))
    return str(file_path)



@pytest.fixture
def banking_pdf(tmp_path):
    """Create a realistic banking document in .pdf format."""
    file_path = tmp_path / "rtgs_spec.pdf"
    
    buffer = BytesIO()
    c = canvas.Canvas(buffer, pagesize=letter)
    
    # Page 1
    c.setFont("Helvetica-Bold", 16)
    c.drawString(100, 750, "RTGS Payment System Specification")
    c.setFont("Helvetica", 12)
    c.drawString(100, 720, "Real Time Gross Settlement System")
    c.drawString(100, 690, "")
    c.drawString(100, 660, "1. Overview")
    c.drawString(100, 630, "RTGS enables real-time transfer of funds between banks.")
    c.drawString(100, 600, "Minimum transaction: Rs. 2,00,000")
    c.drawString(100, 570, "")
    c.drawString(100, 540, "2. Key Features")
    c.drawString(100, 510, "- Real-time settlement")
    c.drawString(100, 480, "- High-value transactions")
    c.drawString(100, 450, "- Immediate confirmation")
    
    c.showPage()
    
    # Page 2
    c.setFont("Helvetica", 12)
    c.drawString(100, 750, "3. Integration Points")
    c.drawString(100, 720, "RTGS integrates with:")
    c.drawString(100, 690, "- Reserve Bank of India (RBI) Gateway")
    c.drawString(100, 660, "- Core Banking System")
    c.drawString(100, 630, "- Transaction Monitoring System")
    
    c.save()
    
    with open(file_path, "wb") as f:
        f.write(buffer.getvalue())
    
    return str(file_path)


class TestBankingDocumentParsing:
    """Test parsing realistic banking documents."""
    
    def test_parse_banking_docx(self, parser, banking_docx):
        """Test parsing a realistic banking .docx document."""
        result = parser.parse(banking_docx, "docx")
        
        # Verify basic structure
        assert result.title == "NEFT Payment System Specification"
        assert len(result.sections) >= 3
        
        # Verify hierarchy is preserved
        levels = [s.level for s in result.sections]
        assert 2 in levels  # Level 2 headings
        assert 3 in levels  # Level 3 headings
        
        # Verify content is extracted
        all_text = " ".join(s.text for s in result.sections)
        assert "NEFT" in all_text
        assert "Core Banking System" in all_text
        assert "2,00,000" in all_text
    
    def test_parse_banking_pdf(self, parser, banking_pdf):
        """Test parsing a realistic banking .pdf document."""
        result = parser.parse(banking_pdf, "pdf")
        
        # Verify basic structure
        assert "RTGS" in result.title or "rtgs" in result.doc_id
        assert len(result.sections) >= 2  # At least 2 pages
        
        # Verify content is extracted
        all_text = " ".join(s.text for s in result.sections)
        assert "RTGS" in all_text
        assert "Real Time" in all_text or "real-time" in all_text.lower()
    
    def test_docx_section_content_accuracy(self, parser, banking_docx):
        """Test that section content is accurately extracted."""
        result = parser.parse(banking_docx, "docx")
        
        # Find the dependencies section
        deps_section = None
        for section in result.sections:
            if "Dependencies" in section.heading or "dependencies" in section.text.lower():
                deps_section = section
                break
        
        assert deps_section is not None
        assert "Core Banking System" in deps_section.text
    
    def test_metadata_completeness(self, parser, banking_docx):
        """Test that all required metadata is present."""
        result = parser.parse(banking_docx, "docx")
        
        required_fields = ["file_name", "file_type", "num_sections"]
        for field in required_fields:
            assert field in result.metadata
        
        assert result.metadata["file_type"] == "docx"
        assert result.metadata["num_sections"] > 0


class TestErrorScenarios:
    """Test error handling with realistic scenarios."""
    
    def test_mixed_content_docx(self, parser, tmp_path):
        """Test .docx with mixed content types."""
        file_path = tmp_path / "mixed.docx"
        doc = Document()
        
        # Mix of headings, paragraphs, and empty lines
        doc.add_heading("Title", level=1)
        doc.add_paragraph("")  # Empty
        doc.add_paragraph("Content")
        doc.add_heading("Section", level=2)
        doc.add_paragraph("")  # Empty
        doc.add_paragraph("More content")
        
        doc.save(str(file_path))
        
        result = parser.parse(str(file_path), "docx")
        assert isinstance(result.sections, list)
        assert len(result.sections) > 0
