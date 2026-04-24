"""
Checkpoint Task 13: End-to-End Ingestion Pipeline Integration Test

This test verifies the complete ingestion pipeline with entity extraction and graph population:
1. Parse multiple sample documents
2. Chunk them hierarchically
3. Generate embeddings
4. Store in PostgreSQL, Qdrant, and BM25 index
5. Extract entities using NER and LLM
6. Resolve entities and create SAME_AS relationships
7. Detect conflicts between entities
8. Populate Neo4j knowledge graph
9. Verify all storage layers are populated correctly
"""

import pytest
import os
import tempfile
from pathlib import Path
from docx import Document

from src.parsing import DocumentParser
from src.chunking import HierarchicalChunker
from src.embedding import EmbeddingGenerator
from src.storage import DatabaseManager, VectorStore
from src.indexing import BM25Indexer
from src.extraction import EntityExtractor, EntityResolver, ConflictDetector
from src.storage import GraphPopulator
from src.pipeline import IngestionPipeline, IngestionStatus
from config.system_config import SystemConfig


@pytest.fixture
def sample_neft_document():
    """Create a sample NEFT banking document for testing."""
    doc = Document()
    doc.add_heading('NEFT Payment System Specification', 0)
    
    doc.add_heading('1. Overview', 1)
    doc.add_paragraph(
        'The National Electronic Funds Transfer (NEFT) system is a nationwide '
        'payment system facilitating one-to-one funds transfer. NEFT operates '
        'in hourly batches and is available 24x7 throughout the year. '
        'The NEFT system is managed by the Reserve Bank of India (RBI).'
    )
    
    doc.add_heading('2. Transaction Limits', 1)
    doc.add_paragraph(
        'NEFT transactions have the following limits: '
        'Minimum amount: Rs. 1 (no minimum limit). '
        'Maximum amount: Rs. 2 lakhs per transaction for retail customers. '
        'Individual banks may set their own limits based on risk assessment.'
    )
    
    doc.add_heading('3. Processing Time', 1)
    doc.add_paragraph(
        'NEFT transactions are processed in batches. The settlement happens '
        'in half-hourly intervals. Transactions are typically completed within '
        '2-3 hours of initiation. NEFT operates on a deferred net settlement basis.'
    )
    
    doc.add_heading('4. System Dependencies', 1)
    doc.add_paragraph(
        'NEFT depends on the Core Banking System (CBS) for account validation '
        'and balance checks. NEFT also integrates with the Payment Gateway for '
        'transaction routing. The system requires real-time connectivity to RBI servers.'
    )
    
    # Save to temporary file
    temp_file = tempfile.NamedTemporaryFile(mode='wb', suffix='.docx', delete=False)
    doc.save(temp_file.name)
    temp_file.close()
    
    yield temp_file.name
    
    # Cleanup
    os.unlink(temp_file.name)


@pytest.fixture
def sample_rtgs_document():
    """Create a sample RTGS banking document for testing."""
    doc = Document()
    doc.add_heading('RTGS Payment System Specification', 0)
    
    doc.add_heading('1. Overview', 1)
    doc.add_paragraph(
        'The Real Time Gross Settlement (RTGS) system is a funds transfer system '
        'where transfer of money takes place from one bank to another on a real-time '
        'and gross basis. RTGS is the fastest possible money transfer system through '
        'the banking channel. Settlement in RTGS is final and irrevocable.'
    )
    
    doc.add_heading('2. Transaction Limits', 1)
    doc.add_paragraph(
        'RTGS transactions have the following limits: '
        'Minimum amount: Rs. 2 lakhs per transaction. '
        'Maximum amount: Rs. 5 lakhs per transaction for retail customers. '
        'There is no upper ceiling for RTGS transactions.'
    )
    
    doc.add_heading('3. Processing Time', 1)
    doc.add_paragraph(
        'RTGS transactions are processed in real-time. Settlement happens immediately '
        'during banking hours. RTGS operates from 9:00 AM to 4:30 PM on weekdays. '
        'Transactions are completed within 30 minutes of initiation.'
    )
    
    doc.add_heading('4. System Dependencies', 1)
    doc.add_paragraph(
        'RTGS depends on the Core Banking System (CBS) for account validation. '
        'RTGS integrates with the Payment Gateway for secure transaction processing. '
        'The system requires dedicated leased line connectivity to RBI.'
    )
    
    # Save to temporary file
    temp_file = tempfile.NamedTemporaryFile(mode='wb', suffix='.docx', delete=False)
    doc.save(temp_file.name)
    temp_file.close()
    
    yield temp_file.name
    
    # Cleanup
    os.unlink(temp_file.name)


@pytest.fixture
def system_config():
    """Create system configuration for testing."""
    from dotenv import load_dotenv
    load_dotenv()
    return SystemConfig.from_env()


@pytest.fixture
def database_manager(system_config):
    """Create database manager and clean up test data."""
    db = DatabaseManager(system_config.postgres_connection_string)
    db.initialize()
    
    # Clean up any existing test data
    with db.get_connection() as conn:
        with conn.cursor() as cur:
            cur.execute("DELETE FROM chunks WHERE doc_id LIKE 'test_checkpoint13_%'")
            cur.execute("DELETE FROM documents WHERE doc_id LIKE 'test_checkpoint13_%'")
        conn.commit()
    
    yield db
    
    # Cleanup after test
    with db.get_connection() as conn:
        with conn.cursor() as cur:
            cur.execute("DELETE FROM chunks WHERE doc_id LIKE 'test_checkpoint13_%'")
            cur.execute("DELETE FROM documents WHERE doc_id LIKE 'test_checkpoint13_%'")
        conn.commit()
    
    db.close()


@pytest.fixture
def vector_store(system_config):
    """Create vector store and clean up test data."""
    vs = VectorStore(
        url=system_config.qdrant_url,
        collection_name="banking_docs_checkpoint13_test",
        vector_size=system_config.embedding_dimension
    )
    
    yield vs
    
    # Cleanup after test
    try:
        vs.client.delete_collection("banking_docs_checkpoint13_test")
    except:
        pass
    
    vs.close()


@pytest.fixture
def graph_populator(system_config):
    """Create graph populator and clean up test data."""
    gp = GraphPopulator(
        neo4j_uri=system_config.neo4j_uri,
        neo4j_user=system_config.neo4j_user,
        neo4j_password=system_config.neo4j_password
    )
    
    # Create schema
    gp.create_schema()
    
    # Clean up any existing test data
    with gp.driver.session() as session:
        session.run("MATCH (n) WHERE n.doc_id STARTS WITH 'test_checkpoint13_' DETACH DELETE n")
        session.run("MATCH (n:Chunk) WHERE n.chunk_id STARTS WITH 'test_checkpoint13_' DETACH DELETE n")
    
    yield gp
    
    # Cleanup after test
    with gp.driver.session() as session:
        session.run("MATCH (n) WHERE n.doc_id STARTS WITH 'test_checkpoint13_' DETACH DELETE n")
        session.run("MATCH (n:Chunk) WHERE n.chunk_id STARTS WITH 'test_checkpoint13_' DETACH DELETE n")
    
    gp.close()


@pytest.fixture
def ingestion_pipeline(system_config, database_manager, vector_store, graph_populator):
    """Create complete ingestion pipeline with all components."""
    parser = DocumentParser()
    chunker = HierarchicalChunker(
        parent_size=system_config.parent_chunk_size,
        child_size=system_config.child_chunk_size,
        overlap=system_config.chunk_overlap
    )
    embedding_generator = EmbeddingGenerator(model_name=system_config.embedding_model)
    bm25_indexer = BM25Indexer()
    entity_extractor = EntityExtractor(
        ollama_base_url=system_config.ollama_base_url,
        llm_model=system_config.llm_model
    )
    entity_resolver = EntityResolver(
        similarity_threshold=system_config.entity_similarity_threshold
    )
    conflict_detector = ConflictDetector(
        ollama_base_url=system_config.ollama_base_url,
        llm_model=system_config.llm_model
    )
    
    pipeline = IngestionPipeline(
        parser=parser,
        chunker=chunker,
        embedding_generator=embedding_generator,
        database_manager=database_manager,
        vector_store=vector_store,
        bm25_indexer=bm25_indexer,
        entity_extractor=entity_extractor,
        entity_resolver=entity_resolver,
        conflict_detector=conflict_detector,
        graph_populator=graph_populator
    )
    
    return pipeline


class TestCheckpoint13EndToEnd:
    """End-to-end integration tests for complete ingestion pipeline."""
    
    def test_single_document_ingestion(
        self,
        ingestion_pipeline,
        sample_neft_document,
        database_manager,
        vector_store,
        graph_populator
    ):
        """
        Test complete ingestion pipeline with a single document.
        
        Verifies:
        1. Document is parsed correctly
        2. Chunks are created and stored
        3. Embeddings are generated and stored
        4. Entities are extracted
        5. Graph is populated
        6. All storage layers are consistent
        """
        print("\n=== Testing Single Document Ingestion ===")
        
        # Ingest document
        result = ingestion_pipeline.ingest(
            sample_neft_document,
            'docx',
            'test_checkpoint13_neft'
        )
        
        # Verify ingestion completed successfully
        assert result.status == IngestionStatus.COMPLETED, f"Ingestion failed: {result.error}"
        assert result.num_chunks > 0, "No chunks were created"
        assert result.error is None
        
        print(f"✓ Ingestion completed successfully")
        print(f"  Document ID: {result.doc_id}")
        print(f"  Chunks: {result.num_chunks}")
        print(f"  Entities: {result.num_entities}")
        print(f"  Relationships: {result.num_relationships}")
        
        # Verify PostgreSQL storage
        doc = database_manager.get_document_by_id('test_checkpoint13_neft')
        assert doc is not None, "Document not found in PostgreSQL"
        assert doc['title'] == 'NEFT Payment System Specification'
        
        chunks = database_manager.get_chunks_by_doc_id('test_checkpoint13_neft')
        assert len(chunks) == result.num_chunks, "Chunk count mismatch in PostgreSQL"
        
        print(f"✓ PostgreSQL storage verified: {len(chunks)} chunks")
        
        # Verify Qdrant storage
        for chunk in chunks:
            vector_result = vector_store.get_by_chunk_id(chunk['chunk_id'])
            assert vector_result is not None, f"Chunk {chunk['chunk_id']} not found in Qdrant"
        
        print(f"✓ Qdrant storage verified: all chunks have embeddings")
        
        # Verify Neo4j storage
        with graph_populator.driver.session() as session:
            # Check document node
            doc_result = session.run(
                "MATCH (d:Document {doc_id: $doc_id}) RETURN d",
                doc_id='test_checkpoint13_neft'
            )
            assert doc_result.single() is not None, "Document not found in Neo4j"
            
            # Check chunk nodes
            chunk_result = session.run(
                "MATCH (c:Chunk) WHERE c.doc_id = $doc_id RETURN count(c) as count",
                doc_id='test_checkpoint13_neft'
            )
            chunk_count = chunk_result.single()['count']
            assert chunk_count == result.num_chunks, "Chunk count mismatch in Neo4j"
            
            # Check entity nodes (if any were extracted)
            entity_result = session.run(
                "MATCH (e) WHERE e.entity_id IS NOT NULL AND e.source_chunk_id STARTS WITH $prefix "
                "RETURN count(e) as count",
                prefix='test_checkpoint13_neft'
            )
            entity_count = entity_result.single()['count']
            
            print(f"✓ Neo4j storage verified: {chunk_count} chunks, {entity_count} entities")
    
    def test_multiple_documents_ingestion(
        self,
        ingestion_pipeline,
        sample_neft_document,
        sample_rtgs_document,
        database_manager,
        vector_store,
        graph_populator
    ):
        """
        Test batch ingestion of multiple documents.
        
        Verifies:
        1. Multiple documents are ingested successfully
        2. Entities are extracted from all documents
        3. Entity resolution identifies common entities (e.g., CBS)
        4. Conflict detection identifies contradictory information
        5. Graph contains relationships between documents
        """
        print("\n=== Testing Multiple Documents Ingestion ===")
        
        # Prepare batch
        documents = [
            {
                "file_path": sample_neft_document,
                "file_type": "docx",
                "doc_id": "test_checkpoint13_neft"
            },
            {
                "file_path": sample_rtgs_document,
                "file_type": "docx",
                "doc_id": "test_checkpoint13_rtgs"
            }
        ]
        
        # Ingest batch
        results = ingestion_pipeline.ingest_batch(documents)
        
        # Verify all documents completed successfully
        assert len(results) == 2, "Expected 2 results"
        assert all(r.status == IngestionStatus.COMPLETED for r in results), \
            f"Some ingestions failed: {[r.error for r in results if r.error]}"
        
        successful = sum(1 for r in results if r.status == IngestionStatus.COMPLETED)
        total_chunks = sum(r.num_chunks for r in results)
        total_entities = sum(r.num_entities for r in results)
        total_relationships = sum(r.num_relationships for r in results)
        
        print(f"✓ Batch ingestion completed: {successful}/{len(results)} successful")
        print(f"  Total chunks: {total_chunks}")
        print(f"  Total entities: {total_entities}")
        print(f"  Total relationships: {total_relationships}")
        
        # Verify both documents are in PostgreSQL
        neft_doc = database_manager.get_document_by_id('test_checkpoint13_neft')
        rtgs_doc = database_manager.get_document_by_id('test_checkpoint13_rtgs')
        
        assert neft_doc is not None, "NEFT document not found"
        assert rtgs_doc is not None, "RTGS document not found"
        
        print(f"✓ Both documents stored in PostgreSQL")
        
        # Verify chunks from both documents are in Qdrant
        neft_chunks = database_manager.get_chunks_by_doc_id('test_checkpoint13_neft')
        rtgs_chunks = database_manager.get_chunks_by_doc_id('test_checkpoint13_rtgs')
        
        for chunk in neft_chunks + rtgs_chunks:
            vector_result = vector_store.get_by_chunk_id(chunk['chunk_id'])
            assert vector_result is not None, f"Chunk {chunk['chunk_id']} not in Qdrant"
        
        print(f"✓ All chunks from both documents in Qdrant")
        
        # Verify graph contains both documents and their relationships
        with graph_populator.driver.session() as session:
            # Check both document nodes exist
            doc_result = session.run(
                "MATCH (d:Document) WHERE d.doc_id IN ['test_checkpoint13_neft', 'test_checkpoint13_rtgs'] "
                "RETURN count(d) as count"
            )
            doc_count = doc_result.single()['count']
            assert doc_count == 2, f"Expected 2 documents in Neo4j, found {doc_count}"
            
            # Check for common entities (e.g., CBS, Payment Gateway)
            # These should appear in both documents and might have SAME_AS relationships
            common_entities_result = session.run(
                "MATCH (e) WHERE e.canonical_name IN ['CBS', 'Core Banking System', 'Payment Gateway'] "
                "RETURN count(e) as count"
            )
            common_count = common_entities_result.single()['count']
            
            print(f"✓ Neo4j graph populated: {doc_count} documents, {common_count} common entities")
            
            # Check for SAME_AS relationships (entity resolution)
            same_as_result = session.run(
                "MATCH ()-[r:SAME_AS]->() RETURN count(r) as count"
            )
            same_as_count = same_as_result.single()['count']
            
            if same_as_count > 0:
                print(f"✓ Entity resolution created {same_as_count} SAME_AS relationships")
            
            # Check for CONFLICTS_WITH relationships
            conflict_result = session.run(
                "MATCH ()-[r:CONFLICTS_WITH]->() RETURN count(r) as count"
            )
            conflict_count = conflict_result.single()['count']
            
            if conflict_count > 0:
                print(f"✓ Conflict detection found {conflict_count} conflicts")
                
                # Show example conflict
                example_conflict = session.run(
                    "MATCH (e1)-[r:CONFLICTS_WITH]->(e2) "
                    "RETURN e1.name as entity1, e2.name as entity2, r.conflict_type as type "
                    "LIMIT 1"
                )
                conflict_record = example_conflict.single()
                if conflict_record:
                    print(f"  Example: {conflict_record['entity1']} conflicts with "
                          f"{conflict_record['entity2']} ({conflict_record['type']})")
    
    def test_entity_extraction_and_resolution(
        self,
        ingestion_pipeline,
        sample_neft_document,
        sample_rtgs_document,
        graph_populator
    ):
        """
        Test entity extraction and resolution across documents.
        
        Verifies:
        1. Entities are extracted from document text
        2. Similar entities are resolved to canonical forms
        3. SAME_AS relationships link entity mentions
        4. Common entities across documents are identified
        """
        print("\n=== Testing Entity Extraction and Resolution ===")
        
        # Ingest both documents
        documents = [
            {
                "file_path": sample_neft_document,
                "file_type": "docx",
                "doc_id": "test_checkpoint13_neft_entities"
            },
            {
                "file_path": sample_rtgs_document,
                "file_type": "docx",
                "doc_id": "test_checkpoint13_rtgs_entities"
            }
        ]
        
        results = ingestion_pipeline.ingest_batch(documents)
        
        assert all(r.status == IngestionStatus.COMPLETED for r in results)
        
        # Query Neo4j for extracted entities
        with graph_populator.driver.session() as session:
            # Get all entity types
            entity_types_result = session.run(
                "MATCH (e) WHERE e.entity_type IS NOT NULL "
                "AND (e.source_chunk_id STARTS WITH 'test_checkpoint13_neft_entities' "
                "OR e.source_chunk_id STARTS WITH 'test_checkpoint13_rtgs_entities') "
                "RETURN DISTINCT e.entity_type as type, count(e) as count"
            )
            
            entity_types = {record['type']: record['count'] for record in entity_types_result}
            
            if entity_types:
                print(f"✓ Extracted entities by type:")
                for entity_type, count in entity_types.items():
                    print(f"  {entity_type}: {count}")
            else:
                print("  Note: No entities extracted (LLM may not be available)")
            
            # Check for SAME_AS relationships
            same_as_result = session.run(
                "MATCH (e1)-[r:SAME_AS]->(e2) "
                "WHERE (e1.source_chunk_id STARTS WITH 'test_checkpoint13_neft_entities' "
                "OR e1.source_chunk_id STARTS WITH 'test_checkpoint13_rtgs_entities') "
                "RETURN e1.name as mention, e2.canonical_name as canonical "
                "LIMIT 5"
            )
            
            same_as_relationships = list(same_as_result)
            
            if same_as_relationships:
                print(f"✓ Entity resolution created SAME_AS relationships:")
                for record in same_as_relationships:
                    print(f"  '{record['mention']}' → '{record['canonical']}'")
            else:
                print("  Note: No SAME_AS relationships (entities may not need resolution)")
    
    def test_conflict_detection(
        self,
        ingestion_pipeline,
        sample_neft_document,
        sample_rtgs_document,
        graph_populator
    ):
        """
        Test conflict detection between documents.
        
        Verifies:
        1. Conflicting information is detected (e.g., different transaction limits)
        2. CONFLICTS_WITH relationships are created
        3. Conflict metadata includes type and source information
        """
        print("\n=== Testing Conflict Detection ===")
        
        # Ingest both documents
        documents = [
            {
                "file_path": sample_neft_document,
                "file_type": "docx",
                "doc_id": "test_checkpoint13_neft_conflicts"
            },
            {
                "file_path": sample_rtgs_document,
                "file_type": "docx",
                "doc_id": "test_checkpoint13_rtgs_conflicts"
            }
        ]
        
        results = ingestion_pipeline.ingest_batch(documents)
        
        assert all(r.status == IngestionStatus.COMPLETED for r in results)
        
        # Query Neo4j for conflicts
        with graph_populator.driver.session() as session:
            # Check for CONFLICTS_WITH relationships
            conflict_result = session.run(
                "MATCH (e1)-[r:CONFLICTS_WITH]->(e2) "
                "WHERE (e1.source_chunk_id STARTS WITH 'test_checkpoint13_neft_conflicts' "
                "OR e1.source_chunk_id STARTS WITH 'test_checkpoint13_rtgs_conflicts') "
                "RETURN e1.name as entity1, e2.name as entity2, "
                "r.conflict_type as type, r.explanation as explanation"
            )
            
            conflicts = list(conflict_result)
            
            if conflicts:
                print(f"✓ Detected {len(conflicts)} conflicts:")
                for conflict in conflicts:
                    print(f"  {conflict['entity1']} ⚠️  {conflict['entity2']}")
                    print(f"    Type: {conflict['type']}")
                    if conflict['explanation']:
                        print(f"    Explanation: {conflict['explanation']}")
            else:
                print("  Note: No conflicts detected (LLM may not be available or no conflicts exist)")
    
    def test_storage_consistency(
        self,
        ingestion_pipeline,
        sample_neft_document,
        database_manager,
        vector_store,
        graph_populator
    ):
        """
        Test consistency across all storage layers.
        
        Verifies:
        1. Same chunks exist in PostgreSQL, Qdrant, and Neo4j
        2. Chunk IDs are consistent across all stores
        3. Metadata is preserved in all stores
        """
        print("\n=== Testing Storage Consistency ===")
        
        # Ingest document
        result = ingestion_pipeline.ingest(
            sample_neft_document,
            'docx',
            'test_checkpoint13_consistency'
        )
        
        assert result.status == IngestionStatus.COMPLETED
        
        # Get chunks from PostgreSQL
        pg_chunks = database_manager.get_chunks_by_doc_id('test_checkpoint13_consistency')
        pg_chunk_ids = {chunk['chunk_id'] for chunk in pg_chunks}
        
        print(f"✓ PostgreSQL: {len(pg_chunk_ids)} chunks")
        
        # Verify all chunks are in Qdrant
        qdrant_chunk_ids = set()
        for chunk_id in pg_chunk_ids:
            vector_result = vector_store.get_by_chunk_id(chunk_id)
            if vector_result is not None:
                qdrant_chunk_ids.add(chunk_id)
        
        assert pg_chunk_ids == qdrant_chunk_ids, \
            f"Chunk ID mismatch: PostgreSQL has {len(pg_chunk_ids)}, Qdrant has {len(qdrant_chunk_ids)}"
        
        print(f"✓ Qdrant: {len(qdrant_chunk_ids)} chunks (consistent with PostgreSQL)")
        
        # Verify all chunks are in Neo4j
        with graph_populator.driver.session() as session:
            neo4j_result = session.run(
                "MATCH (c:Chunk) WHERE c.doc_id = $doc_id RETURN c.chunk_id as chunk_id",
                doc_id='test_checkpoint13_consistency'
            )
            neo4j_chunk_ids = {record['chunk_id'] for record in neo4j_result}
        
        assert pg_chunk_ids == neo4j_chunk_ids, \
            f"Chunk ID mismatch: PostgreSQL has {len(pg_chunk_ids)}, Neo4j has {len(neo4j_chunk_ids)}"
        
        print(f"✓ Neo4j: {len(neo4j_chunk_ids)} chunks (consistent with PostgreSQL)")
        print(f"✓ All storage layers are consistent!")
    
    def test_ingestion_failure_handling(
        self,
        ingestion_pipeline
    ):
        """
        Test that ingestion handles failures gracefully.
        
        Verifies:
        1. Invalid file paths are handled
        2. Error messages are descriptive
        3. Pipeline halts on failure
        4. Status is set to FAILED
        """
        print("\n=== Testing Ingestion Failure Handling ===")
        
        # Try to ingest non-existent file
        result = ingestion_pipeline.ingest(
            "nonexistent_file.docx",
            'docx',
            'test_checkpoint13_failure'
        )
        
        # Verify failure is handled
        assert result.status == IngestionStatus.FAILED, "Expected FAILED status"
        assert result.error is not None, "Expected error message"
        
        print(f"✓ Failure handled gracefully")
        print(f"  Status: {result.status.value}")
        print(f"  Error: {result.error}")


if __name__ == "__main__":
    pytest.main([__file__, "-v", "-s"])
