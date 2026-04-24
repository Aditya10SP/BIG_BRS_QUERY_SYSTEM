"""Unit tests for DocumentParser."""

import pytest
from pathlib import Path
from docx import Document
from docx.shared import Pt
from io import BytesIO
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter

from src.parsing import DocumentParser, ParsedDocument, Section, ParsingError


@pytest.fixture
def parser():
    """Create DocumentParser instance."""
    return DocumentParser()


@pytest.fixture
def temp_docx_file(tmp_path):
    """Create a temporary .docx file for testing."""
    def _create_docx(filename="test.docx", content=None):
        file_path = tmp_path / filename
        doc = Document()
        
        if content is None:
            # Default content with structure
            doc.add_heading("Test Document Title", level=1)
            doc.add_paragraph("This is the introduction paragraph.")
            doc.add_heading("Section 1", level=2)
            doc.add_paragraph("Content for section 1.")
            doc.add_heading("Section 2", level=2)
            doc.add_paragraph("Content for section 2.")
        else:
            # Custom content
            for item in content:
                if item["type"] == "heading":
                    doc.add_heading(item["text"], level=item.get("level", 1))
                elif item["type"] == "paragraph":
                    doc.add_paragraph(item["text"])
        
        doc.save(str(file_path))
        return str(file_path)
    
    return _create_docx


@pytest.fixture
def temp_pdf_file(tmp_path):
    """Create a temporary .pdf file for testing."""
    def _create_pdf(filename="test.pdf", content=None):
        file_path = tmp_path / filename
        
        # Create PDF with reportlab
        buffer = BytesIO()
        c = canvas.Canvas(buffer, pagesize=letter)
        
        if content is None:
            # Default content
            c.drawString(100, 750, "Test PDF Document")
            c.drawString(100, 700, "This is page 1 content.")
            c.showPage()
            c.drawString(100, 750, "Page 2 content here.")
            c.showPage()
        else:
            # Custom content
            y_position = 750
            for text in content:
                c.drawString(100, y_position, text)
                y_position -= 20
                if y_position < 100:
                    c.showPage()
                    y_position = 750
        
        c.save()
        
        # Write to file
        with open(file_path, "wb") as f:
            f.write(buffer.getvalue())
        
        return str(file_path)
    
    return _create_pdf


class TestDocumentParserBasics:
    """Test basic DocumentParser functionality."""
    
    def test_parser_initialization(self, parser):
        """Test parser can be initialized."""
        assert parser is not None
        assert parser.SUPPORTED_FORMATS == {".docx", ".pdf"}
    
    def test_unsupported_file_type(self, parser, tmp_path):
        """Test error on unsupported file type."""
        file_path = tmp_path / "test.txt"
        file_path.write_text("test content")
        
        with pytest.raises(ParsingError) as exc_info:
            parser.parse(str(file_path), "txt")
        
        assert "Unsupported file format" in str(exc_info.value)
    
    def test_file_not_found(self, parser):
        """Test error when file doesn't exist."""
        with pytest.raises(ParsingError) as exc_info:
            parser.parse("/nonexistent/file.docx", "docx")
        
        assert "File not found" in str(exc_info.value)
    
    def test_path_is_directory(self, parser, tmp_path):
        """Test error when path is a directory."""
        with pytest.raises(ParsingError) as exc_info:
            parser.parse(str(tmp_path), "docx")
        
        assert "not a file" in str(exc_info.value)


class TestDocxParsing:
    """Test .docx file parsing."""
    
    def test_parse_simple_docx(self, parser, temp_docx_file):
        """Test parsing a simple .docx file."""
        file_path = temp_docx_file()
        result = parser.parse(file_path, "docx")
        
        assert isinstance(result, ParsedDocument)
        assert result.doc_id == "test"
        assert result.title == "Test Document Title"
        assert len(result.sections) > 0
        assert result.metadata["file_type"] == "docx"
    
    def test_docx_sections_extracted(self, parser, temp_docx_file):
        """Test that sections are properly extracted from .docx."""
        file_path = temp_docx_file()
        result = parser.parse(file_path, "docx")
        
        # Should have sections for the headings
        assert len(result.sections) >= 2
        
        # Check section structure
        for section in result.sections:
            assert isinstance(section, Section)
            assert section.section_id
            assert section.heading
            assert section.level > 0
            assert isinstance(section.text, str)
    
    def test_docx_hierarchy_preserved(self, parser, temp_docx_file):
        """Test that heading hierarchy is preserved."""
        content = [
            {"type": "heading", "text": "Main Title", "level": 1},
            {"type": "paragraph", "text": "Intro text"},
            {"type": "heading", "text": "Subsection A", "level": 2},
            {"type": "paragraph", "text": "Content A"},
            {"type": "heading", "text": "Subsection B", "level": 2},
            {"type": "paragraph", "text": "Content B"},
            {"type": "heading", "text": "Sub-subsection", "level": 3},
            {"type": "paragraph", "text": "Nested content"},
        ]
        file_path = temp_docx_file(content=content)
        result = parser.parse(file_path, "docx")
        
        # Check levels are preserved
        levels = [s.level for s in result.sections]
        assert 1 in levels
        assert 2 in levels
        assert 3 in levels
    
    def test_docx_metadata_extracted(self, parser, temp_docx_file):
        """Test that metadata is extracted from .docx."""
        file_path = temp_docx_file()
        result = parser.parse(file_path, "docx")
        
        assert "file_name" in result.metadata
        assert "file_type" in result.metadata
        assert "num_paragraphs" in result.metadata
        assert "num_sections" in result.metadata
        assert result.metadata["file_name"] == "test.docx"
    
    def test_docx_empty_document(self, parser, tmp_path):
        """Test parsing empty .docx document."""
        file_path = tmp_path / "empty.docx"
        doc = Document()
        doc.save(str(file_path))
        
        result = parser.parse(str(file_path), "docx")
        
        # Should still return a valid ParsedDocument
        assert isinstance(result, ParsedDocument)
        # May have empty or default sections
        assert isinstance(result.sections, list)
    
    def test_docx_no_headings(self, parser, temp_docx_file):
        """Test .docx with only paragraphs, no headings."""
        content = [
            {"type": "paragraph", "text": "First paragraph"},
            {"type": "paragraph", "text": "Second paragraph"},
            {"type": "paragraph", "text": "Third paragraph"},
        ]
        file_path = temp_docx_file(content=content)
        result = parser.parse(file_path, "docx")
        
        # Should create at least one section
        assert len(result.sections) >= 1
        assert result.sections[0].text
    
    def test_corrupted_docx(self, parser, tmp_path):
        """Test error handling for corrupted .docx file."""
        file_path = tmp_path / "corrupted.docx"
        file_path.write_bytes(b"This is not a valid docx file")
        
        with pytest.raises(ParsingError) as exc_info:
            parser.parse(str(file_path), "docx")
        
        assert "Corrupted or invalid" in str(exc_info.value)


class TestPdfParsing:
    """Test .pdf file parsing."""
    
    def test_parse_simple_pdf(self, parser, temp_pdf_file):
        """Test parsing a simple .pdf file."""
        file_path = temp_pdf_file()
        result = parser.parse(file_path, "pdf")
        
        assert isinstance(result, ParsedDocument)
        assert result.doc_id == "test"
        assert result.title  # Should have some title
        assert len(result.sections) > 0
        assert result.metadata["file_type"] == "pdf"
    
    def test_pdf_sections_per_page(self, parser, temp_pdf_file):
        """Test that PDF creates sections per page."""
        file_path = temp_pdf_file()
        result = parser.parse(file_path, "pdf")
        
        # Should have sections for each page
        assert len(result.sections) >= 2
        
        # Check page numbers are tracked
        for section in result.sections:
            assert len(section.page_numbers) > 0
    
    def test_pdf_text_extraction(self, parser, temp_pdf_file):
        """Test that text is extracted from PDF."""
        content = ["Line 1 of text", "Line 2 of text", "Line 3 of text"]
        file_path = temp_pdf_file(content=content)
        result = parser.parse(file_path, "pdf")
        
        # Check that text was extracted
        all_text = " ".join(s.text for s in result.sections)
        assert len(all_text) > 0
    
    def test_pdf_metadata_extracted(self, parser, temp_pdf_file):
        """Test that metadata is extracted from .pdf."""
        file_path = temp_pdf_file()
        result = parser.parse(file_path, "pdf")
        
        assert "file_name" in result.metadata
        assert "file_type" in result.metadata
        assert "num_pages" in result.metadata
        assert "num_sections" in result.metadata
        assert result.metadata["file_name"] == "test.pdf"
    
    def test_corrupted_pdf(self, parser, tmp_path):
        """Test error handling for corrupted .pdf file."""
        file_path = tmp_path / "corrupted.pdf"
        file_path.write_bytes(b"This is not a valid pdf file")
        
        with pytest.raises(ParsingError) as exc_info:
            parser.parse(str(file_path), "pdf")
        
        assert "Corrupted or invalid" in str(exc_info.value)


class TestCaseInsensitivity:
    """Test case-insensitive file type handling."""
    
    def test_uppercase_docx(self, parser, temp_docx_file):
        """Test parsing with uppercase file type."""
        file_path = temp_docx_file()
        result = parser.parse(file_path, "DOCX")
        assert isinstance(result, ParsedDocument)
    
    def test_uppercase_pdf(self, parser, temp_pdf_file):
        """Test parsing with uppercase file type."""
        file_path = temp_pdf_file()
        result = parser.parse(file_path, "PDF")
        assert isinstance(result, ParsedDocument)
    
    def test_mixed_case(self, parser, temp_docx_file):
        """Test parsing with mixed case file type."""
        file_path = temp_docx_file()
        result = parser.parse(file_path, "DocX")
        assert isinstance(result, ParsedDocument)


class TestDataclasses:
    """Test dataclass structures."""
    
    def test_section_dataclass(self):
        """Test Section dataclass."""
        section = Section(
            section_id="s1",
            heading="Test Section",
            level=1,
            text="Test content",
            page_numbers=[1, 2]
        )
        assert section.section_id == "s1"
        assert section.heading == "Test Section"
        assert section.level == 1
        assert section.text == "Test content"
        assert section.page_numbers == [1, 2]
    
    def test_section_default_page_numbers(self):
        """Test Section with default page_numbers."""
        section = Section(
            section_id="s1",
            heading="Test",
            level=1,
            text="Content"
        )
        assert section.page_numbers == []
    
    def test_parsed_document_dataclass(self):
        """Test ParsedDocument dataclass."""
        sections = [
            Section("s1", "Heading", 1, "Text", [1])
        ]
        doc = ParsedDocument(
            doc_id="doc1",
            title="Test Doc",
            sections=sections,
            metadata={"key": "value"}
        )
        assert doc.doc_id == "doc1"
        assert doc.title == "Test Doc"
        assert len(doc.sections) == 1
        assert doc.metadata == {"key": "value"}
    
    def test_parsed_document_default_metadata(self):
        """Test ParsedDocument with default metadata."""
        doc = ParsedDocument(
            doc_id="doc1",
            title="Test",
            sections=[]
        )
        assert doc.metadata == {}
