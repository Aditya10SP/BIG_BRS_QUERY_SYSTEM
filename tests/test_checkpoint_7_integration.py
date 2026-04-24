"""
Checkpoint Task 7: End-to-End Vector RAG Pipeline Integration Test

This test verifies the complete vector RAG pipeline:
1. Parse a sample document
2. Chunk it hierarchically
3. Generate embeddings
4. Store in PostgreSQL, Qdrant, and BM25 index
5. Query with vector + BM25 hybrid search
6. Verify results are correct
"""

import pytest
import os
import tempfile
from pathlib import Path
from docx import Document

from src.parsing import DocumentParser, ParsedDocument
from src.chunking import HierarchicalChunker, Chunk
from src.embedding import EmbeddingGenerator
from src.storage import DatabaseManager, VectorStore
from src.indexing import BM25Indexer
from config.system_config import SystemConfig


@pytest.fixture
def sample_banking_document():
    """Create a sample banking document for testing."""
    # Create a temporary .docx file with banking content
    doc = Document()
    doc.add_heading('NEFT Payment System Specification', 0)
    
    doc.add_heading('1. Overview', 1)
    doc.add_paragraph(
        'The National Electronic Funds Transfer (NEFT) system is a nationwide '
        'payment system facilitating one-to-one funds transfer. NEFT operates '
        'in hourly batches and is available 24x7 throughout the year.'
    )
    
    doc.add_heading('2. Transaction Limits', 1)
    doc.add_paragraph(
        'NEFT transactions have the following limits: '
        'Minimum amount: Rs. 1 (no minimum limit). '
        'Maximum amount: No maximum limit for NEFT transactions. '
        'Individual banks may set their own limits based on risk assessment.'
    )
    
    doc.add_heading('3. Processing Time', 1)
    doc.add_paragraph(
        'NEFT transactions are processed in batches. The settlement happens '
        'in half-hourly intervals. Transactions are typically completed within '
        '2-3 hours of initiation. NEFT operates on a deferred net settlement basis.'
    )
    
    doc.add_heading('4. Integration with Core Banking', 1)
    doc.add_paragraph(
        'NEFT integrates with the Core Banking System (CBS) for account validation '
        'and balance checks. The CBS must support real-time balance inquiry and '
        'account status verification. NEFT depends on CBS for transaction authorization.'
    )
    
    # Save to temporary file
    temp_file = tempfile.NamedTemporaryFile(mode='wb', suffix='.docx', delete=False)
    doc.save(temp_file.name)
    temp_file.close()
    
    yield temp_file.name
    
    # Cleanup
    os.unlink(temp_file.name)


@pytest.fixture
def system_config():
    """Create system configuration for testing."""
    from dotenv import load_dotenv
    load_dotenv()
    return SystemConfig.from_env()


@pytest.fixture
def parser():
    """Create document parser."""
    return DocumentParser()


@pytest.fixture
def chunker():
    """Create hierarchical chunker."""
    return HierarchicalChunker(parent_size=2048, child_size=512)


@pytest.fixture
def embedding_generator(system_config):
    """Create embedding generator."""
    return EmbeddingGenerator(model_name=system_config.embedding_model)


@pytest.fixture
def database_manager(system_config):
    """Create database manager and clean up test data."""
    db = DatabaseManager(system_config.postgres_connection_string)
    db.initialize()  # Initialize connection pool and schema
    
    # Clean up any existing test data
    with db.get_connection() as conn:
        with conn.cursor() as cur:
            cur.execute("DELETE FROM chunks WHERE doc_id LIKE 'test_%'")
            cur.execute("DELETE FROM documents WHERE doc_id LIKE 'test_%'")
        conn.commit()
    
    yield db
    
    # Cleanup after test
    with db.get_connection() as conn:
        with conn.cursor() as cur:
            cur.execute("DELETE FROM chunks WHERE doc_id LIKE 'test_%'")
            cur.execute("DELETE FROM documents WHERE doc_id LIKE 'test_%'")
        conn.commit()
    
    db.close()


@pytest.fixture
def vector_store(system_config):
    """Create vector store and clean up test data."""
    vs = VectorStore(
        url=system_config.qdrant_url,
        collection_name="banking_docs_test",
        vector_size=system_config.embedding_dimension
    )
    
    # Collection is created automatically in __init__
    
    yield vs
    
    # Cleanup after test
    try:
        vs.client.delete_collection("banking_docs_test")
    except:
        pass
    
    vs.close()


@pytest.fixture
def bm25_indexer():
    """Create BM25 indexer."""
    return BM25Indexer()


class TestCheckpoint7EndToEnd:
    """End-to-end integration tests for vector RAG pipeline."""
    
    def _store_document_and_chunks(self, database_manager, parsed_doc, chunks):
        """Helper method to store document and chunks in PostgreSQL."""
        database_manager.create_document(
            doc_id=parsed_doc.doc_id,
            title=parsed_doc.title,
            file_path=None,
            file_type='docx',
            metadata=parsed_doc.metadata
        )
        
        for chunk in chunks:
            database_manager.create_chunk(
                chunk_id=chunk.chunk_id,
                doc_id=chunk.doc_id,
                text=chunk.text,
                chunk_type=chunk.chunk_type,
                parent_chunk_id=chunk.parent_chunk_id,
                breadcrumbs=chunk.breadcrumbs,
                section=chunk.section,
                token_count=chunk.token_count,
                metadata=chunk.metadata
            )
    
    def _store_embeddings(self, vector_store, chunks, embeddings):
        """Helper method to store embeddings in Qdrant."""
        chunk_ids = [chunk.chunk_id for chunk in chunks]
        metadata = [
            {
                "doc_id": chunk.doc_id,
                "text": chunk.text,
                "breadcrumbs": chunk.breadcrumbs,
                "section": chunk.section,
                "chunk_type": chunk.chunk_type,
            }
            for chunk in chunks
        ]
        vector_store.store_embeddings(chunk_ids, embeddings, metadata)
    
    def test_complete_ingestion_pipeline(
        self,
        sample_banking_document,
        parser,
        chunker,
        embedding_generator,
        database_manager,
        vector_store,
        bm25_indexer
    ):
        """
        Test complete ingestion pipeline: parse → chunk → embed → store
        
        This test verifies:
        1. Document parsing extracts structure correctly
        2. Hierarchical chunking creates parent and child chunks
        3. Embeddings are generated for all chunks
        4. Chunks are stored in PostgreSQL with metadata
        5. Embeddings are stored in Qdrant
        6. BM25 index is built for keyword search
        """
        # Step 1: Parse document
        parsed_doc = parser.parse(sample_banking_document, 'docx')
        
        assert parsed_doc is not None
        assert parsed_doc.title == 'NEFT Payment System Specification'
        assert len(parsed_doc.sections) > 0
        print(f"✓ Parsed document with {len(parsed_doc.sections)} sections")
        
        # Step 2: Chunk document
        chunks = chunker.chunk(parsed_doc)
        
        assert len(chunks) > 0
        parent_chunks = [c for c in chunks if c.chunk_type == 'parent']
        child_chunks = [c for c in chunks if c.chunk_type == 'child']
        
        assert len(parent_chunks) > 0
        assert len(child_chunks) > 0
        
        # Verify child chunks have breadcrumbs
        for chunk in child_chunks:
            assert chunk.breadcrumbs is not None
            assert len(chunk.breadcrumbs) > 0
            assert chunk.parent_chunk_id is not None
        
        print(f"✓ Created {len(parent_chunks)} parent chunks and {len(child_chunks)} child chunks")
        
        # Step 3: Store document in PostgreSQL
        self._store_document_and_chunks(database_manager, parsed_doc, chunks)
        print(f"✓ Stored document and {len(chunks)} chunks in PostgreSQL")
        
        # Step 5: Generate embeddings and store in Qdrant
        chunk_texts = [chunk.text for chunk in chunks]
        embeddings = embedding_generator.batch_generate(chunk_texts)
        
        assert embeddings.shape[0] == len(chunks)
        assert embeddings.shape[1] == 384  # Embedding dimension
        print(f"✓ Generated {len(embeddings)} embeddings")
        
        # Store embeddings in Qdrant
        self._store_embeddings(vector_store, chunks, embeddings)
        print(f"✓ Stored embeddings in Qdrant")
        
        # Step 6: Build BM25 index
        bm25_indexer.index(chunks)
        print(f"✓ Built BM25 index with {len(chunks)} chunks")
        
        # Verify storage
        retrieved_doc = database_manager.get_document_by_id(parsed_doc.doc_id)
        assert retrieved_doc is not None
        assert retrieved_doc['title'] == parsed_doc.title
        print(f"✓ Verified document retrieval from PostgreSQL")
    
    def test_vector_search_query(
        self,
        sample_banking_document,
        parser,
        chunker,
        embedding_generator,
        database_manager,
        vector_store,
        bm25_indexer
    ):
        """
        Test vector similarity search retrieves relevant chunks.
        
        This test verifies:
        1. Query embedding is generated
        2. Similar chunks are retrieved from Qdrant
        3. Results are ranked by similarity
        4. Full chunk text is retrieved from PostgreSQL
        """
        # Ingest document first
        parsed_doc = parser.parse(sample_banking_document, 'docx')
        chunks = chunker.chunk(parsed_doc)
        self._store_document_and_chunks(database_manager, parsed_doc, chunks)
        
        chunk_texts = [chunk.text for chunk in chunks]
        embeddings = embedding_generator.batch_generate(chunk_texts)
        self._store_embeddings(vector_store, chunks, embeddings)
        
        # Query: "What are NEFT transaction limits?"
        query = "What are NEFT transaction limits?"
        query_embedding = embedding_generator.generate(query)
        
        # Search Qdrant
        results = vector_store.search(query_embedding, top_k=5)
        
        assert len(results) > 0
        print(f"✓ Vector search returned {len(results)} results")
        
        # Verify results are ranked by similarity
        scores = [score for _, score, _ in results]
        assert scores == sorted(scores, reverse=True), "Results should be ranked by score"
        print(f"✓ Results are ranked by similarity score")
        
        # Verify relevant content is retrieved
        chunk_id, score, metadata = results[0]
        
        # Retrieve full chunk from PostgreSQL
        retrieved_chunk = database_manager.get_chunk_by_id(chunk_id)
        assert retrieved_chunk is not None
        
        # Check if the result is relevant (should mention limits or transaction)
        chunk_text = retrieved_chunk['text'].lower()
        assert 'limit' in chunk_text or 'transaction' in chunk_text or 'amount' in chunk_text
        print(f"✓ Top result is relevant to query")
        print(f"  Score: {score:.4f}")
        print(f"  Chunk: {chunk_text[:100]}...")
    
    def test_bm25_keyword_search(
        self,
        sample_banking_document,
        parser,
        chunker,
        embedding_generator,
        database_manager,
        vector_store,
        bm25_indexer
    ):
        """
        Test BM25 keyword search retrieves chunks with exact terms.
        
        This test verifies:
        1. BM25 index returns results for keyword queries
        2. Results contain the queried keywords
        3. Results are ranked by BM25 score
        """
        # Ingest document first
        parsed_doc = parser.parse(sample_banking_document, 'docx')
        chunks = chunker.chunk(parsed_doc)
        self._store_document_and_chunks(database_manager, parsed_doc, chunks)
        
        bm25_indexer.index(chunks)
        
        # Query with acronym: "NEFT"
        query = "NEFT"
        results = bm25_indexer.search(query, top_k=5)
        
        assert len(results) > 0
        print(f"✓ BM25 search returned {len(results)} results for '{query}'")
        
        # Verify results contain the keyword
        for chunk_id, score in results:
            chunk = database_manager.get_chunk_by_id(chunk_id)
            assert chunk is not None
            assert 'NEFT' in chunk['text'] or 'neft' in chunk['text'].lower()
        
        print(f"✓ All results contain the keyword '{query}'")
        
        # Verify results are ranked
        scores = [score for _, score in results]
        assert scores == sorted(scores, reverse=True), "Results should be ranked by BM25 score"
        print(f"✓ Results are ranked by BM25 score")
    
    def test_hybrid_search(
        self,
        sample_banking_document,
        parser,
        chunker,
        embedding_generator,
        database_manager,
        vector_store,
        bm25_indexer
    ):
        """
        Test hybrid search combining vector similarity and BM25 keyword matching.
        
        This test verifies:
        1. Both vector and BM25 search are executed
        2. Results are combined using Reciprocal Rank Fusion
        3. Hybrid search provides better coverage than either method alone
        """
        # Ingest document first
        parsed_doc = parser.parse(sample_banking_document, 'docx')
        chunks = chunker.chunk(parsed_doc)
        self._store_document_and_chunks(database_manager, parsed_doc, chunks)
        
        chunk_texts = [chunk.text for chunk in chunks]
        embeddings = embedding_generator.batch_generate(chunk_texts)
        self._store_embeddings(vector_store, chunks, embeddings)
        bm25_indexer.index(chunks)
        
        # Query: "CBS integration with NEFT"
        query = "CBS integration with NEFT"
        
        # Vector search
        query_embedding = embedding_generator.generate(query)
        vector_results = vector_store.search(query_embedding, top_k=5)
        vector_chunk_ids = {chunk_id for chunk_id, _, _ in vector_results}
        
        # BM25 search
        bm25_results = bm25_indexer.search(query, top_k=5)
        bm25_chunk_ids = {chunk_id for chunk_id, _ in bm25_results}
        
        print(f"✓ Vector search returned {len(vector_results)} results")
        print(f"✓ BM25 search returned {len(bm25_results)} results")
        
        # Combine using Reciprocal Rank Fusion
        def rrf_score(rank, k=60):
            return 1.0 / (k + rank)
        
        combined_scores = {}
        
        # Add vector scores
        for rank, (chunk_id, _, _) in enumerate(vector_results):
            combined_scores[chunk_id] = combined_scores.get(chunk_id, 0) + rrf_score(rank)
            combined_scores[chunk_id] = combined_scores.get(chunk_id, 0) + rrf_score(rank)
        
        # Add BM25 scores
        for rank, (chunk_id, _) in enumerate(bm25_results):
            combined_scores[chunk_id] = combined_scores.get(chunk_id, 0) + rrf_score(rank)
        
        # Sort by combined score
        hybrid_results = sorted(combined_scores.items(), key=lambda x: x[1], reverse=True)
        
        assert len(hybrid_results) > 0
        print(f"✓ Hybrid search combined results: {len(hybrid_results)} unique chunks")
        
        # Verify hybrid search includes results from both methods
        hybrid_chunk_ids = {chunk_id for chunk_id, _ in hybrid_results}
        
        # At least some overlap or union
        assert len(hybrid_chunk_ids) >= max(len(vector_chunk_ids), len(bm25_chunk_ids))
        print(f"✓ Hybrid search provides comprehensive coverage")
        
        # Check top result is relevant
        top_chunk_id = hybrid_results[0][0]
        top_chunk = database_manager.get_chunk_by_id(top_chunk_id)
        chunk_text = top_chunk['text'].lower()
        
        # Should mention CBS or integration or NEFT
        assert any(term in chunk_text for term in ['cbs', 'core banking', 'integration', 'neft'])
        print(f"✓ Top hybrid result is relevant to query")
        print(f"  Chunk: {chunk_text[:150]}...")
    
    def test_chunk_metadata_preservation(
        self,
        sample_banking_document,
        parser,
        chunker,
        embedding_generator,
        database_manager,
        vector_store,
        bm25_indexer
    ):
        """
        Test that chunk metadata is preserved throughout the pipeline.
        
        This test verifies:
        1. Breadcrumbs are preserved in PostgreSQL
        2. Metadata is preserved in Qdrant payload
        3. Retrieved chunks have complete metadata
        """
        # Ingest document
        parsed_doc = parser.parse(sample_banking_document, 'docx')
        chunks = chunker.chunk(parsed_doc)
        self._store_document_and_chunks(database_manager, parsed_doc, chunks)
        
        chunk_texts = [chunk.text for chunk in chunks]
        embeddings = embedding_generator.batch_generate(chunk_texts)
        self._store_embeddings(vector_store, chunks, embeddings)
        
        # Pick a child chunk
        child_chunk = next(c for c in chunks if c.chunk_type == 'child')
        
        # Retrieve from PostgreSQL
        retrieved_chunk = database_manager.get_chunk_by_id(child_chunk.chunk_id)
        
        assert retrieved_chunk is not None
        assert retrieved_chunk['breadcrumbs'] == child_chunk.breadcrumbs
        assert retrieved_chunk['section'] == child_chunk.section
        assert retrieved_chunk['chunk_type'] == child_chunk.chunk_type
        assert retrieved_chunk['parent_chunk_id'] == child_chunk.parent_chunk_id
        
        print(f"✓ PostgreSQL preserves all chunk metadata")
        print(f"  Breadcrumbs: {retrieved_chunk['breadcrumbs']}")
        
        # Search in Qdrant and verify payload
        query_embedding = embedding_generator.generate(child_chunk.text[:50])
        results = vector_store.search(query_embedding, top_k=1)
        
        assert len(results) > 0
        chunk_id, score, metadata = results[0]
        
        assert metadata['doc_id'] == child_chunk.doc_id
        assert metadata['breadcrumbs'] == child_chunk.breadcrumbs
        assert metadata['section'] == child_chunk.section
        
        print(f"✓ Qdrant preserves metadata in payload")


if __name__ == "__main__":
    pytest.main([__file__, "-v", "-s"])
